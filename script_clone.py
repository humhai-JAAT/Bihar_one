import streamlit as st
import pandas as pd
from docx import Document
import base64
from pathlib import Path
import sys
sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

# Page configuration
st.set_page_config(page_title="GRCS Simulator", layout="wide", initial_sidebar_state="collapsed")

original_st_dataframe = st.dataframe


def render_centered_dataframe(data, **kwargs):
    if isinstance(data, pd.DataFrame):
        styled_df = data.style.set_properties(**{"text-align": "center"}).set_table_styles(
            [
                {"selector": "th", "props": [("text-align", "center")]},
                {"selector": "td", "props": [("text-align", "center")]},
            ]
        )
        return original_st_dataframe(styled_df, **kwargs)
    return original_st_dataframe(data, **kwargs)


st.dataframe = render_centered_dataframe

# Define GRCS Reference Data (Global - used across Simulator, Weight, and Reference pages)
reference_data = [
    {"S.No": 1, "Attribute": "Aadhaar", "Weight (%)": 7.111274871, "Match Type": "Deterministic", "Enterprise Rule": "UIDAI biometric verified"},
    {"S.No": 2, "Attribute": "Name", "Weight (%)": 4.605747973, "Match Type": "Fuzzy + Phonetic", "Enterprise Rule": "UIDAI > Civil Registry precedence"},
    {"S.No": 3, "Attribute": "Date of Birth", "Weight (%)": 5.600589536, "Match Type": "Exact > Year", "Enterprise Rule": "Civil Registry override"},
    {"S.No": 4, "Attribute": "Mobile Number", "Weight (%)": 4.532056006, "Match Type": "OTP Verified", "Enterprise Rule": "Aadhaar seeded + CBS timestamp"},
    {"S.No": 5, "Attribute": "Gender", "Weight (%)": 4.679439941, "Match Type": "Exact", "Enterprise Rule": "Legal identity anchor"},
    {"S.No": 6, "Attribute": "Father's Name", "Weight (%)": 5.011053795, "Match Type": "Fuzzy", "Enterprise Rule": "Civil Registry priority"},
    {"S.No": 7, "Attribute": "Mother's Name", "Weight (%)": 5.011053795, "Match Type": "Fuzzy", "Enterprise Rule": "Civil Registry validated"},
    {"S.No": 8, "Attribute": "Permanent Address", "Weight (%)": 4.56890199, "Match Type": "Geo-normalized", "Enterprise Rule": "UIDAI > Land Registry"},
    {"S.No": 9, "Attribute": "Correspondence Address", "Weight (%)": 3.463522476, "Match Type": "Latest Timestamp", "Enterprise Rule": "CBS latest update"},
    {"S.No": 10, "Attribute": "Caste", "Weight (%)": 5.342667649, "Match Type": "Certificate Verified", "Enterprise Rule": "RTPS validated"},
    {"S.No": 11, "Attribute": "Marital Status", "Weight (%)": 4.237288136, "Match Type": "Registry Preferred", "Enterprise Rule": "Marriage Registry > Self"},
    {"S.No": 12, "Attribute": "Education Status", "Weight (%)": 4.016212233, "Match Type": "Dept Certified", "Enterprise Rule": "Education DB"},
    {"S.No": 13, "Attribute": "Employment Status", "Weight (%)": 3.831982314, "Match Type": "Statutory", "Enterprise Rule": "Labour Dept verified"},
    {"S.No": 14, "Attribute": "Ration Card Number", "Weight (%)": 5.490051584, "Match Type": "Deterministic", "Enterprise Rule": "PDS Household anchor"},
    {"S.No": 15, "Attribute": "Ration Card Type", "Weight (%)": 4.089042, "Match Type": "Exact", "Enterprise Rule": "Welfare classification"},
    {"S.No": 16, "Attribute": "PAN ID", "Weight (%)": 6.632277082, "Match Type": "Deterministic", "Enterprise Rule": "Income Tax authority"},
    {"S.No": 17, "Attribute": "Bank Account", "Weight (%)": 5.711127487, "Match Type": "Masked Deterministic", "Enterprise Rule": "CBS source-of-origin"},
    {"S.No": 18, "Attribute": "Land Ownership", "Weight (%)": 6.042741341, "Match Type": "Legal Title", "Enterprise Rule": "Land Registry override"},
    {"S.No": 19, "Attribute": "Motor Ownership", "Weight (%)": 5.416359617, "Match Type": "Registration Match", "Enterprise Rule": "VAHAN verified"},
    {"S.No": 20, "Attribute": "Nationality", "Weight (%)": 4.605747973, "Match Type": "Legal", "Enterprise Rule": "Civil Registry"}
]

# Custom CSS for modern UI with top navbar
st.markdown("""
<style>
    html, body {
        margin: 0 !important;
        padding: 0 !important;
    }

    /* Hide Streamlit header and toolbar */
    header {display: none !important;}
    [data-testid="stHeader"] {height: 0 !important;}
    [data-testid="stHeader"] {display: none !important;}
    [data-testid="stToolbar"] {display: none !important;}
    [data-testid="stDecoration"] {display: none !important;}
    #MainMenu {display: none !important;}
    .stDeployButton {display: none !important;}
    footer {display: none !important;}

    [data-testid="stAppViewContainer"] {
        padding-top: 0 !important;
        margin-top: 0 !important;
    }

    [data-testid="stAppViewContainer"] > .main {
        padding-top: 0 !important;
        margin-top: 0 !important;
    }
    
    /* Hide default sidebar */
    [data-testid="collapsedControl"] {display: none;}
    
    /* Light Blue Background */
    .stApp {
        background: linear-gradient(180deg, #e3f2fd 0%, #bbdefb 100%);
    }
    
    /* Main container styling */
    .main {
        padding-top: 0rem !important;
        background: transparent;
    }
    .block-container {
        padding-top: 3.5rem !important;
        background: transparent;
    }
    
    /* Typography - Improved Visibility */
    h1 {font-size: 32px !important; font-weight: 700 !important; color: #0d47a1; margin-bottom: 1rem;}
    h2 {font-size: 22px !important; font-weight: 600 !important; color: #1565c0; margin-top: 1.5rem; margin-bottom: 1rem;}
    h3 {font-size: 16px !important; font-weight: 600 !important; color: #1976d2; margin-top: 1rem; margin-bottom: 0.5rem;}
    p, label, div {font-size: 15px !important; color: #263238; font-weight: 500;}

    /* Text alignment */
    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] li,
    [data-testid="stMarkdownContainer"] span,
    .stText,
    .stCaption {
        text-align: justify !important;
    }

    /* Table alignment */
    [data-testid="stTable"] table th,
    [data-testid="stTable"] table td {
        text-align: center !important;
    }

    [data-testid="stDataFrame"] th,
    [data-testid="stDataFrame"] td,
    [data-testid="stDataFrame"] [role="columnheader"],
    [data-testid="stDataFrame"] [role="gridcell"] {
        text-align: center !important;
        justify-content: center !important;
    }

    /* Override Streamlit default dark text shades */
    [data-testid="stAppViewContainer"] {
        --text-color: #0d47a1;
    }
    .stApp [style*="color:#FFFFFF"],
    .stApp [style*="color: #262730"],
    .stApp [style*="color:rgb(38,39,48)"],
    .stApp [style*="color: rgb(38, 39, 48)"] {
        color: #0d47a1 !important;
    }
    
    /* Centered Page Header */
    .page-header {
        text-align: center;
        padding: 1.5rem 1.5rem;
        background: white;
        border-radius: 12px;
        box-shadow: 0 4px 16px rgba(13, 71, 161, 0.12);
        margin: 0.2rem auto 1.5rem auto;
        max-width: 900px;
        border: 2px solid #2196f3;
    }
    .page-header h1 {
        font-size: 28px !important;
        color: #0d47a1 !important;
        margin-bottom: 0.5rem !important;
        font-weight: 800 !important;
    }
    .page-header p {
        font-size: 13px !important;
        color: #37474f !important;
        margin: 0 !important;
        font-weight: 400 !important;
    }
    
    /* Streamlit components - Better Visibility */
    .stSlider, .stSelectbox, .stCheckbox {
        font-size: 14px !important;
        color: #263238 !important;
        font-weight: 600 !important;
    }

    /* Selectbox label visibility */
    .stSelectbox label {
        color: #000000 !important;
        font-weight: 600 !important;
    }

    /* Selectbox readability (Risk Level and similar dropdowns) */
    .stSelectbox [data-baseweb="select"] > div {
        background-color: #FFFFFF !important;
        color: #000000 !important;
        border: 2px solid #999999 !important;
    }

    .stSelectbox [data-baseweb="select"] input,
    .stSelectbox [data-baseweb="select"] span,
    .stSelectbox [data-baseweb="select"] div {
        color: #000000 !important;
        background-color: #FFFFFF !important;
    }

    div[role="listbox"] {
        background-color: #FFFFFF !important;
        border: 2px solid #999999 !important;
    }

    div[role="option"] {
        background-color: #FFFFFF !important;
        color: #000000 !important;
        font-weight: 500 !important;
    }

    div[role="option"][aria-selected="true"],
    div[role="option"]:hover {
        background-color: #FFFFFF !important;
        color: #000000 !important;
        font-weight: 700 !important;
        border-left: 3px solid #0d47a1 !important;
    }

    /* Additional selectbox styling */
    .stSelectbox div[data-baseweb="select"] {
        background-color: #FFFFFF !important;
    }

    .stSelectbox div[data-baseweb="select"] > div {
        background-color: #FFFFFF !important;
    }

    /* Comprehensive dropdown styling (popover layer) */
    div[data-baseweb="popover"],
    div[data-baseweb="popover"] > div,
    div[data-baseweb="popover"] [role="listbox"],
    div[data-baseweb="popover"] ul,
    div[data-baseweb="popover"] li {
        background-color: #FFFFFF !important;
        color: #0d2238 !important;
        border-color: #90caf9 !important;
    }

    div[data-baseweb="popover"] [role="option"],
    div[data-baseweb="popover"] [role="option"] *,
    div[data-baseweb="popover"] li,
    div[data-baseweb="popover"] li * {
        background-color: #FFFFFF !important;
        color: #0d2238 !important;
        opacity: 1 !important;
    }

    div[data-baseweb="popover"] [role="option"]:hover,
    div[data-baseweb="popover"] [role="option"][aria-selected="true"],
    div[data-baseweb="popover"] li:hover,
    div[data-baseweb="popover"] li[aria-selected="true"] {
        background-color: #E3F2FD !important;
        color: #0d2238 !important;
        border-left: 3px solid #2196f3 !important;
    }
    .stButton>button {
        background: linear-gradient(135deg, #90caf9 0%, #64b5f6 100%);
        color: #0d47a1;
        border: none;
        padding: 0.6rem 2rem;
        border-radius: 8px;
        font-weight: 700;
        font-size: 14px;
        transition: all 0.3s;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(100, 181, 246, 0.45);
        background: linear-gradient(135deg, #bbdefb 0%, #90caf9 100%);
    }
    
    /* Expander / dropdown bar styling (readable + consistent) */
    [data-testid="stExpander"] {
        border: 1px solid #90caf9 !important;
        border-radius: 10px !important;
        overflow: hidden !important;
        background: #ffffff !important;
    }

    [data-testid="stExpander"] [data-testid="stExpanderHeader"],
    [data-testid="stExpander"] summary {
        background: #e3f2fd !important;
        color: #0d47a1 !important;
        border: 0 !important;
    }

    [data-testid="stExpander"] [data-testid="stExpanderHeader"] *,
    [data-testid="stExpander"] summary * {
        color: #0d47a1 !important;
        font-weight: 700 !important;
        font-size: 15px !important;
    }

    [data-testid="stExpander"] summary:hover,
    [data-testid="stExpander"] [data-testid="stExpanderHeader"]:hover {
        background: #d2e8fb !important;
    }

    [data-testid="stExpander"] summary:focus,
    [data-testid="stExpander"] summary:focus-visible,
    [data-testid="stExpander"] [data-testid="stExpanderHeader"]:focus,
    [data-testid="stExpander"] [data-testid="stExpanderHeader"]:focus-visible {
        background: #d2e8fb !important;
        color: #0d47a1 !important;
        outline: 2px solid #64b5f6 !important;
        outline-offset: -2px !important;
    }

    [data-testid="stExpander"] svg {
        fill: #0d47a1 !important;
        color: #0d47a1 !important;
    }
    
    /* Dataframe styling */
    .dataframe {
        border: 2px solid #2196f3 !important;
        border-radius: 8px;
        background: white !important;
    }
    
    /* Metric styling */
    [data-testid="stMetricValue"] {
        font-size: 26px !important;
        color: #0d47a1 !important;
        font-weight: 800 !important;
    }
    
    /* Custom cards - Blue Theme */
    .info-card {
        background: linear-gradient(135deg, #90caf9 0%, #64b5f6 100%);
        padding: 1.5rem;
        border-radius: 12px;
        color: #0d47a1;
        margin-bottom: 1rem;
        box-shadow: 0 6px 20px rgba(100, 181, 246, 0.35);
        text-align: center;
    }

    .info-card h1,
    .info-card p {
        color: #0d47a1 !important;
        text-align: center !important;
    }
    
    .result-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        border-left: 5px solid #64b5f6;
        box-shadow: 0 4px 16px rgba(100, 181, 246, 0.25);
        margin: 1rem 0;
    }
    
    /* Input labels with better visibility */
    label {
        color: #0d47a1 !important;
        font-weight: 700 !important;
        font-size: 14px !important;
    }

</style>
""", unsafe_allow_html=True)

# Top Navbar with Logos and Navigation
st.markdown("""
<style>
    .top-navbar {
        background: linear-gradient(135deg, #bbdefb 0%, #90caf9 100%);
        padding: 1rem 2rem;
        display: flex;
        align-items: center;
        justify-content: space-between;
        width: 100vw;
        margin: 0;
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        z-index: 1000;
        box-shadow: 0 4px 20px rgba(0,0,0,0.1);
        border-radius: 0;
    }
    
    .navbar-logos {
        display: flex;
        align-items: center;
        gap: 2rem;
    }
    
    .navbar-logo {
        height: 50px;
        width: auto;
        background: white;
        padding: 8px;
        border-radius: 8px;
    }

    .navbar-left,
    .navbar-right {
        display: flex;
        align-items: center;
    }
    
    .navbar-title {
        color: #0d47a1;
        font-size: 24px;
        font-weight: 700;
        margin: 0 2rem;
        flex-grow: 1;
    }
    
    .nav-tabs {
        display: flex;
        gap: 0.5rem;
        flex-wrap: wrap;
        background: rgba(255,255,255,0.45);
        padding: 0.5rem;
        border-radius: 10px;
    }
    
    .nav-tab {
        background: rgba(255,255,255,0.7);
        color: #0d47a1;
        padding: 0.6rem 1.2rem;
        border-radius: 8px;
        text-decoration: none;
        font-weight: 600;
        font-size: 13px;
        transition: all 0.3s;
        cursor: pointer;
        border: 2px solid transparent;
        white-space: nowrap;
    }
    
    .nav-tab:hover {
        background: white;
        color: #0d47a1;
        transform: translateY(-2px);
    }
    
    .nav-tab.active {
        background: white;
        color: #0d47a1;
        border-color: white;
    }
</style>
""", unsafe_allow_html=True)

# Render top navbar logos from assets
# Use absolute path based on script location for Render deployment compatibility
script_dir = Path(__file__).parent
logo_bihar_path = script_dir / "assets" / "bihargovt-logo.png"
logo_cipl_path = script_dir / "assets" / "cipl-logo.png"

if logo_bihar_path.exists() and logo_cipl_path.exists():
    logo_bihar_b64 = base64.b64encode(logo_bihar_path.read_bytes()).decode()
    logo_cipl_b64 = base64.b64encode(logo_cipl_path.read_bytes()).decode()

    st.markdown(f"""
    <div class="top-navbar">
        <div class="navbar-left">
            <img src="data:image/png;base64,{logo_cipl_b64}" class="navbar-logo" alt="CIPL Logo">
        </div>
        <h2 class="navbar-title">Golden Record Confidence Score (GRCS)</h2>
        <div class="navbar-right">
            <img src="data:image/png;base64,{logo_bihar_b64}" class="navbar-logo" alt="Bihar Government Logo">
        </div>
    </div>
    """, unsafe_allow_html=True)

# Navigation tabs
col1, col2, col3, col4, col5 = st.columns(5)
with col1:
    simulator_btn = st.button("Simulator", use_container_width=True, type="secondary" if st.session_state.get("page", "Simulator") != "Simulator" else "primary")
with col2:
    ref_btn = st.button("GRCS Reference", use_container_width=True, type="secondary" if st.session_state.get("page", "Simulator") != "Reference" else "primary")
with col3:
    doc_btn = st.button("Technical Documentation", use_container_width=True, type="secondary" if st.session_state.get("page", "Simulator") != "Documentation" else "primary")
with col4:
    weight_btn = st.button("Weight Calculation", use_container_width=True, type="secondary" if st.session_state.get("page", "Simulator") != "Weight" else "primary")
with col5:
    lusr_btn = st.button("LUSR Calculation", use_container_width=True, type="secondary" if st.session_state.get("page", "Simulator") != "LUSR" else "primary")

# Handle navigation
if simulator_btn:
    st.session_state.page = "Simulator"
elif ref_btn:
    st.session_state.page = "Reference"
elif doc_btn:
    st.session_state.page = "Documentation"
elif weight_btn:
    st.session_state.page = "Weight"
elif lusr_btn:
    st.session_state.page = "LUSR"

# Initialize default page
if "page" not in st.session_state:
    st.session_state.page = "Simulator"

page = st.session_state.page

st.markdown("---")

if page == "Simulator":
    st.markdown('''
    <div class="page-header">
        <h1 style="color: white; margin: 0;">Golden Record Confidence Score (GRCS) Simulator</h1>
    </div>
    ''', unsafe_allow_html=True)

    # Extract attributes and weights from global reference_data
    attributes = {item["Attribute"]: item["Weight (%)"] for item in reference_data}
    source_authority = {
        "UIDAI": 85,
        "Civil Registry": 80,
        "RTPS Certified": 82,
        "Income Tax": 78,
        "Bank CBS": 78,
        "Self Declared": 20
    }
    st.header("Attribute Matching Section")
    total_score = 0
    results = []
    for attr, weight in attributes.items():
        with st.expander(f"**{attr}** (Weight: {weight})", expanded=True):
            col1, col2 = st.columns(2)

            with col1:
                mi = st.slider(f"Match Strength (Mi) for {attr}", 0.0, 1.0, 1.0, key=f"mi_{attr}")

            with col2:
                source = st.selectbox(f"Source for {attr}", list(source_authority.keys()), key=f"src_{attr}")

            si = source_authority[source] / 100

            contribution = weight * mi * si
            total_score += contribution

            st.write(f"Contribution: {round(contribution,2)}")
            results.append({
                "Attribute": attr,
                "Weight": weight,
                "Mi": mi,
                "Si": si,
                "Contribution": contribution
            })


    # Reinforcement
    st.header("Final GRCS Calculation")
    df = pd.DataFrame(results)

    tab1, tab2 = st.tabs(["Graph", "Table"])
    with tab1:
       st.bar_chart(df.set_index("Attribute")["Contribution"])
    with tab2:
        st.dataframe(df, use_container_width=True)

    reinforcement = 0
    if st.checkbox("Aadhaar + Name + DOB Exact Match"):
        reinforcement = 5
        st.write("Reinforcement Applied: +5")

    # Risk Adjustment
    st.header("Risk Adjustment")

    risk_level = st.selectbox("Risk Level", ["Low", "Medium", "High"])

    risk_factor = 0
    if risk_level == "Medium":
        risk_factor = 0.02
    elif risk_level == "High":
        risk_factor = 0.05

    # Final GRCS
    grcs = total_score + reinforcement
    grcs = grcs * (1 - risk_factor)

    st.header("Final Result")

    st.subheader(f"GRCS Score: {round(grcs,2)}%")

    # Decision Logic
    if grcs >= 92:
        decision = "Auto Merge"
    elif grcs >= 80:
        decision = "Conditional Auto Merge"
    elif grcs >= 70:
        decision = "Steward Assisted Merge"
    elif grcs >= 60:
        decision = "Manual Validation"
    else:
        decision = "Create New Golden Record"

    st.subheader(f"Decision: {decision}")






elif page == "Reference":
    st.markdown('<div class="info-card">'
    '<h1 style="color: white; margin: 0;">GRCS Reference Table</h1>'
    '<p style="color: rgba(255,255,255,0.9); margin-top: 0.5rem;"></div>', unsafe_allow_html=True)
    
    st.markdown("### Complete GRCS Attribute Reference")

    reference_data = [
        {"S.No": 1, "Attribute": "Aadhaar", "Weight (%)": 7.111274871, "Match Type": "Deterministic", "Enterprise Rule": "UIDAI biometric verified"},
        {"S.No": 2, "Attribute": "Name", "Weight (%)": 4.605747973, "Match Type": "Fuzzy + Phonetic", "Enterprise Rule": "UIDAI > Civil Registry precedence"},
        {"S.No": 3, "Attribute": "Date of Birth", "Weight (%)": 5.600589536, "Match Type": "Exact > Year", "Enterprise Rule": "Civil Registry override"},
        {"S.No": 4, "Attribute": "Mobile Number", "Weight (%)": 4.532056006, "Match Type": "OTP Verified", "Enterprise Rule": "Aadhaar seeded + CBS timestamp"},
        {"S.No": 5, "Attribute": "Gender", "Weight (%)": 4.679439941, "Match Type": "Exact", "Enterprise Rule": "Legal identity anchor"},
        {"S.No": 6, "Attribute": "Father's Name", "Weight (%)": 5.011053795, "Match Type": "Fuzzy", "Enterprise Rule": "Civil Registry priority"},
        {"S.No": 7, "Attribute": "Mother's Name", "Weight (%)": 5.011053795, "Match Type": "Fuzzy", "Enterprise Rule": "Civil Registry validated"},
        {"S.No": 8, "Attribute": "Permanent Address", "Weight (%)": 4.56890199, "Match Type": "Geo-normalized", "Enterprise Rule": "UIDAI > Land Registry"},
        {"S.No": 9, "Attribute": "Correspondence Address", "Weight (%)": 3.463522476, "Match Type": "Latest Timestamp", "Enterprise Rule": "CBS latest update"},
        {"S.No": 10, "Attribute": "Caste", "Weight (%)": 5.342667649, "Match Type": "Certificate Verified", "Enterprise Rule": "RTPS validated"},
        {"S.No": 11, "Attribute": "Marital Status", "Weight (%)": 4.237288136, "Match Type": "Registry Preferred", "Enterprise Rule": "Marriage Registry > Self"},
        {"S.No": 12, "Attribute": "Education Status", "Weight (%)": 4.016212233, "Match Type": "Dept Certified", "Enterprise Rule": "Education DB"},
        {"S.No": 13, "Attribute": "Employment Status", "Weight (%)": 3.831982314, "Match Type": "Statutory", "Enterprise Rule": "Labour Dept verified"},
        {"S.No": 14, "Attribute": "Ration Card Number", "Weight (%)": 5.490051584, "Match Type": "Deterministic", "Enterprise Rule": "PDS Household anchor"},
        {"S.No": 15, "Attribute": "Ration Card Type", "Weight (%)": 4.089042, "Match Type": "Exact", "Enterprise Rule": "Welfare classification"},
        {"S.No": 16, "Attribute": "PAN ID", "Weight (%)": 6.632277082, "Match Type": "Deterministic", "Enterprise Rule": "Income Tax authority"},
        {"S.No": 17, "Attribute": "Bank Account", "Weight (%)": 5.711127487, "Match Type": "Masked Deterministic", "Enterprise Rule": "CBS source-of-origin"},
        {"S.No": 18, "Attribute": "Land Ownership", "Weight (%)": 6.042741341, "Match Type": "Legal Title", "Enterprise Rule": "Land Registry override"},
        {"S.No": 19, "Attribute": "Motor Ownership", "Weight (%)": 5.416359617, "Match Type": "Registration Match", "Enterprise Rule": "VAHAN verified"},
        {"S.No": 20, "Attribute": "Nationality", "Weight (%)": 4.605747973, "Match Type": "Legal", "Enterprise Rule": "Civil Registry"}
    ]

    df = pd.DataFrame(reference_data)
    display_df = df.copy()
    display_df["Weight (%)"] = display_df["Weight (%)"].astype(str) + "%"
    
    # Display the table
    st.dataframe(display_df, width='stretch', hide_index=True)
    
    # Show some statistics
    st.markdown("### Key Insights")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Attributes", len(df))
    with col2:
        st.metric("Total Weight", f"{df['Weight (%)'].sum():.2f}%")
    with col3:
        st.metric("Max Weight Attribute", df.loc[df['Weight (%)'].idxmax(), 'Attribute'])
    with col4:
        st.metric("Avg Weight per Attribute", f"{df['Weight (%)'].mean():.2f}%")
    
    # Download option
    st.markdown("### Download Data")
    csv = display_df.to_csv(index=False)
    st.download_button(
        label="Download as CSV",
        data=csv,
        file_name="GRCS_Reference.csv",
        mime="text/csv"
    )

elif page == "Documentation":
    st.markdown('<div class="info-card"><h1 style="color: white; margin: 0;">Technical Documentation</h1><p style="color: rgba(255,255,255,0.9); margin-top: 0.5rem;">Comprehensive GRCS methodology and matching algorithms</p></div>', unsafe_allow_html=True)
    
    # Read the DOCX file
    doc_file = "GRCS_Technical_Documentation.docx"
    doc = Document(doc_file)
    
    # Extract and display document content
    st.markdown("### Complete Technical Documentation")
    
    # Display all paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Style headings
            if text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                st.markdown(f"#### {text}")
                
                # Add Match Strength table after section 5
                if text.startswith("5. Match Strength"):
                    st.markdown("**Match Strength Reference Table:**")
                    match_strength_data = {
                        "Match Scenario": ["Exact Match", "92% Similarity", "Year of Birth Only Matches", "Unverified Mobile"],
                        "Mi Value": [1.0, 0.92, 0.6, 0.5]
                    }
                    df_mi = pd.DataFrame(match_strength_data)
                    st.dataframe(df_mi, width='stretch', hide_index=True)
            else:
                st.markdown(text)
    
    # Display tables from document
    if doc.tables:
        st.markdown("### Reference Tables from Documentation")
        for i, table in enumerate(doc.tables):
            st.markdown(f"#### Reference Table {i+1}")
            
            # Convert table to dataframe
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            
            if data:
                headers = data[0] if len(data) > 0 else []
                rows = data[1:] if len(data) > 1 else []
                df_table = pd.DataFrame(rows, columns=headers)
                st.dataframe(df_table, width='stretch', hide_index=True)
    
    # Append scoring engine documentation
    st.markdown("---")
    st.markdown("## 10. CONFIDENCE & MATCH SCORING ENGINE")
    st.markdown("*Enterprise Production Model – GRCS Framework*")
    
    st.markdown("""
The Citizen Golden Record (CGR) implements a mathematically governed, enterprise-grade scoring engine to determine identity match strength, attribute trust, and final merge decisions.

This framework transforms CGR into a Trust-Indexed Identity Infrastructure, ensuring that every merge, update, and survivorship decision is quantifiable, explainable, and policy-aligned.

The Golden Record Confidence Score (GRCS) is computed using a structured three-layer model:
- **Wi** – Attribute Weight (importance of attribute)
- **Mi** – Match Strength (quality of match)
- **Si** – Source Trust Multiplier (authority strength)

### 10.1 Final Golden Record Confidence Score (GRCS)

The production scoring formula is:

**GRCS = ∑ (Wi × Mi × Si) + DeterministicReinforcement − RiskAdjustment**

Where:
- **Wi** = Attribute Weight (%)
- **Mi** = Match Strength (0–1 scale)
- **Si** = Source Trust Multiplier (0–1 scale)
- **n** = Total attributes evaluated
- **DeterministicReinforcement** = High-assurance identity boost
- **RiskAdjustment** = Risk moderation factor (if applicable)

Since total Wi = 100%, GRCS is naturally expressed as a percentage (0–100%).

GRCS is recalculated on every update event and stored historically for audit and trend analysis.

---

### 10.2 Finalized Attribute Weight Allocation (Derived from ACS Model)
""")
    
    # Weights allocation table
    weights_data = {
        "S.No": list(range(1, 21)),
        "Attribute": ["Aadhaar", "Name", "Date of Birth", "Mobile Number", "Gender", "Father's Name", "Mother's Name", 
                     "Permanent Address", "Correspondence Address", "Caste", "Marital Status", "Education Status", 
                     "Employment Status", "Ration Card Number", "Ration Card Type", "PAN ID", "Bank Account", 
                     "Land Ownership", "Motor Ownership", "Nationality"],
        "Weight (%)": [18, 9, 9, 7, 3, 6, 4, 8, 4, 4, 2, 2, 2, 5, 2, 5, 4, 3, 2, 1],
        "Match Type": ["Deterministic", "Fuzzy + Phonetic", "Exact > Year", "OTP Verified", "Exact", "Fuzzy", "Fuzzy", 
                      "Geo-normalized", "Latest Timestamp", "Certificate Verified", "Registry Preferred", "Dept Certified",
                      "Statutory", "Deterministic", "Exact", "Deterministic", "Masked Deterministic", "Legal Title", 
                      "Registration Match", "Legal"],
        "Enterprise Rule": ["UIDAI biometric verified", "UIDAI > Civil Registry precedence", "Civil Registry override", 
                           "Aadhaar seeded + CBS timestamp", "Legal identity anchor", "Civil Registry priority", 
                           "Civil Registry validated", "UIDAI > Land Registry", "CBS latest update", "RTPS validated",
                           "Marriage Registry > Self", "Education DB", "Labour Dept verified", "PDS Household anchor",
                           "Welfare classification", "Income Tax authority", "CBS source-of-origin", "Land Registry override",
                           "VAHAN verified", "Civil Registry"]
    }
    df_weights = pd.DataFrame(weights_data)
    st.dataframe(df_weights, width='stretch', hide_index=True)
    
    st.markdown("""
These weights are derived using the Attribute Criticality Score (ACS) model and are configurable under State Data Governance oversight.

---

### 10.3 Attribute Weight Derivation Framework (ACS Model)

Weights are derived using:

**ACS_i = (0.35×L_i + 0.30×U_i + 0.20×S_i + 0.15×R_i)**

Where:
- **L_i** = Legal Strength (0–10)
- **U_i** = Uniqueness Power (0–10)
- **S_i** = Stability Over Time (0–10)
- **R_i** = Service Impact Risk (0–10)

**Final Weight:**
**W_i = (ACS_i / ∑ACS) × 100**

This ensures:
- Legally strong attributes receive higher weight
- Highly unique identifiers dominate identity resolution
- Stable attributes carry long-term trust
- Service-critical attributes influence merge decisions

Weights are periodically reviewed by the State Data Governance Council.

---

### 10.4 Match Strength Model (M_i)

Match Strength reflects how accurately an incoming value matches an existing attribute.

**Deterministic Match:**
- Exact match → 1.0
- No match → 0

**Exact Match:**
- Full match → 1.0
- Partial match (e.g., year-only DOB) → 0.6

**Fuzzy Match:**
- Based on Levenshtein / Jaro-Winkler similarity
- 92% similarity → 0.92

**OTP / Timestamp Based:**
- Verified & latest → 1.0
- Verified but old → 0.8
- Unverified → 0.5

---

### 10.5 Source Authority Index (SAI) & Trust Multiplier (S_i)

Each source is assigned an Authority Score (0–100):
""")
    
    # Authority scores table
    authority_data = {
        "Source": ["UIDAI", "Civil Registry", "RTPS Certified Docs", "Income Tax Dept", "Bank CBS", "Land Registry", 
                   "Transport Registry", "PDS", "Survey DB", "Self Declared"],
        "Authority Score": [85, 80, 82, 78, 78, 75, 75, 70, 45, 20]
    }
    df_authority = pd.DataFrame(authority_data)
    st.dataframe(df_authority, width='stretch', hide_index=True)
    
    st.markdown("""
**Trust Multiplier:**
**S_i = AuthorityScore / 100**

If RTPS-certified document exists:

**S_i = min(S_i + 0.05, 1.0)**

Authority scores are governed and periodically recalibrated.

---

### 10.6 Deterministic Reinforcement Factor

A reinforcement factor is applied when multi-anchor identity alignment is detected (e.g., Aadhaar + Name + DOB exact match).

This factor provides proportional confidence enhancement for high-assurance identity alignment and is capped to prevent overweight dominance.

---

### 10.7 Risk Adjustment Layer (Optional)

For high-risk or fraud-sensitive attributes, a moderation factor may be applied:

**AdjustedScore = GRCS × (1 − RiskFactor)**

Where:
- Low Risk → 0
- Medium Risk → 0.02
- High Risk → 0.05

This ensures sensitive attributes are risk-calibrated without destabilizing the scoring framework.

---

### 10.8 Decision Engine Thresholds
""")
    
    # Decision thresholds table
    decision_data = {
        "GRCS (%)": ["≥ 92%", "80–91%", "70–79%", "60–69%", "<60%"],
        "Decision": ["Auto Merge", "Conditional Auto Merge + Audit Log", "Steward Assisted Merge", "Manual Validation", "Create New Golden Record"]
    }
    df_decision = pd.DataFrame(decision_data)
    st.dataframe(df_decision, width='stretch', hide_index=True)
    
    st.markdown("""
Thresholds are configurable under governance control.

---

### 10.9 Conflict Resolution Logic

If attribute conflict occurs:

1. Compare Survivorship Priority Index (SPI)
2. Compare Source Trust Multiplier (S_i)
3. If dynamic field → latest verified timestamp prevails
4. If score difference <10% → steward review

This ensures legally authoritative and higher-trust sources prevail.

---

### 10.10 Explainability & Auditability

Each attribute stores:
- **W_i** (weight)
- **M_i** (match strength)
- **S_i** (trust multiplier)
- Reinforcement applied
- Risk adjustment applied
- Final contribution to GRCS

**Example:**
```json
{
  "attribute": "Aadhaar",
  "weight": 18,
  "match_strength": 1.0,
  "trust_multiplier": 0.85,
  "contribution": 15.3
}
```

This guarantees:
- Full transparency
- Governance auditability
- Legal defensibility
- RTPS compliance

---

### 10.11 Strategic Outcome

The GRCS framework ensures that:
- Identity is anchored on legally authoritative sources
- Matching is mathematically governed
- Weight distribution is policy-driven
- Conflict resolution is structured
- Steward intervention is threshold-based
- Every merge decision is explainable

**The Golden Record is therefore a governed enterprise scoring engine — not a heuristic matching system.**
""")
    
    # Download documentation
    st.markdown("### Download Documentation")
    with open(doc_file, 'rb') as f:
        doc_bytes = f.read()
    st.download_button(
        label="Download DOCX",
        data=doc_bytes,
        file_name="GRCS_Technical_Documentation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

elif page == "Weight":
    st.markdown('<div class="info-card"><h1 style="color: white; margin: 0;">Weight Calculation</h1><p style="color: rgba(255,255,255,0.9); margin-top: 0.5rem;">Calculate attribute weights using ACS model (L, U, S, R parameters)</p></div>', unsafe_allow_html=True)
    
    # ========== WEIGHT CALCULATION CALCULATOR ==========
    st.header("Weight Calculation Calculator (Based on ACS Model)")
    
    st.markdown("""
    ### Formula:
    **ACS_i = (0.35 × L_i + 0.30 × U_i + 0.20 × S_i + 0.15 × R_i)**
    
    **Then:** Wi = (ACS_i / Total ACS) × 100
    
    **Parameters (each 0-10 scale):**
    - **L = Legal Strength** (35%)
    - **U = Uniqueness Power** (30%)
    - **S = Stability Over Time** (20%)
    - **R = Service Impact Risk** (15%)
    """)
    
    st.markdown("---")
    st.markdown("### Enter L, U, S, R Values (0-10) for Each Attribute:")
    
    # Extract attribute names from global reference_data
    attributes_list = [item["Attribute"] for item in reference_data]
    
    # Create LUSR Table 6 data for dynamic slider ranges
    lusr_table_6_data = [
        {"Attribute": "Aadhaar", "Legal Strength (L)": 10, "Uniqueness (U)": 10, "Stability (S)": 9, "Service Risk Impact (R)": 9, "ACL": 9.65},
        {"Attribute": "Name", "Legal Strength (L)": 7, "Uniqueness (U)": 5, "Stability (S)": 7, "Service Risk Impact (R)": 6, "ACL": 6.25},
        {"Attribute": "Date of Birth", "Legal Strength (L)": 8, "Uniqueness (U)": 6, "Stability (S)": 9, "Service Risk Impact (R)": 8, "ACL": 7.6},
        {"Attribute": "Mobile Number", "Legal Strength (L)": 6, "Uniqueness (U)": 6, "Stability (S)": 6, "Service Risk Impact (R)": 7, "ACL": 6.15},
        {"Attribute": "Gender", "Legal Strength (L)": 7, "Uniqueness (U)": 4, "Stability (S)": 9, "Service Risk Impact (R)": 6, "ACL": 6.35},
        {"Attribute": "Father's Name", "Legal Strength (L)": 8, "Uniqueness (U)": 5, "Stability (S)": 8, "Service Risk Impact (R)": 6, "ACL": 6.8},
        {"Attribute": "Mother's Name", "Legal Strength (L)": 8, "Uniqueness (U)": 5, "Stability (S)": 8, "Service Risk Impact (R)": 6, "ACL": 6.8},
        {"Attribute": "Permanent Address", "Legal Strength (L)": 7, "Uniqueness (U)": 5, "Stability (S)": 6, "Service Risk Impact (R)": 7, "ACL": 6.2},
        {"Attribute": "Correspondence Address", "Legal Strength (L)": 5, "Uniqueness (U)": 4, "Stability (S)": 5, "Service Risk Impact (R)": 5, "ACL": 4.7},
        {"Attribute": "Caste", "Legal Strength (L)": 8, "Uniqueness (U)": 5, "Stability (S)": 8, "Service Risk Impact (R)": 9, "ACL": 7.25},
        {"Attribute": "Marital Status", "Legal Strength (L)": 7, "Uniqueness (U)": 4, "Stability (S)": 6, "Service Risk Impact (R)": 6, "ACL": 5.75},
        {"Attribute": "Education Status", "Legal Strength (L)": 6, "Uniqueness (U)": 4, "Stability (S)": 7, "Service Risk Impact (R)": 5, "ACL": 5.45},
        {"Attribute": "Employment Status", "Legal Strength (L)": 6, "Uniqueness (U)": 4, "Stability (S)": 5, "Service Risk Impact (R)": 6, "ACL": 5.2},
        {"Attribute": "Ration Card Number", "Legal Strength (L)": 7, "Uniqueness (U)": 8, "Stability (S)": 7, "Service Risk Impact (R)": 8, "ACL": 7.45},
        {"Attribute": "Ration Card Type", "Legal Strength (L)": 6, "Uniqueness (U)": 4, "Stability (S)": 6, "Service Risk Impact (R)": 7, "ACL": 5.55},
        {"Attribute": "PAN ID", "Legal Strength (L)": 9, "Uniqueness (U)": 9, "Stability (S)": 9, "Service Risk Impact (R)": 9, "ACL": 9.0},
        {"Attribute": "Bank Account", "Legal Strength (L)": 8, "Uniqueness (U)": 8, "Stability (S)": 6, "Service Risk Impact (R)": 9, "ACL": 7.75},
        {"Attribute": "Land Ownership", "Legal Strength (L)": 9, "Uniqueness (U)": 7, "Stability (S)": 8, "Service Risk Impact (R)": 9, "ACL": 8.2},
        {"Attribute": "Motor Ownership", "Legal Strength (L)": 8, "Uniqueness (U)": 7, "Stability (S)": 7, "Service Risk Impact (R)": 7, "ACL": 7.35},
        {"Attribute": "Nationality", "Legal Strength (L)": 8, "Uniqueness (U)": 3, "Stability (S)": 9, "Service Risk Impact (R)": 5, "ACL": 6.25}
    ]
    
    # Create mapping for quick attribute lookup
    lusr_mapping = {item["Attribute"]: item for item in lusr_table_6_data}
    
    # Store ACS calculations
    acs_data = []
    
    for attr in attributes_list:
        with st.expander(f"**{attr}**", expanded=True):
            # Get max values from Table 6 for this attribute
            attr_data = lusr_mapping.get(attr, {"Legal Strength (L)": 10, "Uniqueness (U)": 10, "Stability (S)": 10, "Service Risk Impact (R)": 10})
            max_L = attr_data.get("Legal Strength (L)", 10)
            max_U = attr_data.get("Uniqueness (U)", 10)
            max_S = attr_data.get("Stability (S)", 10)
            max_R = attr_data.get("Service Risk Impact (R)", 10)
            
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                L = st.slider("L", 0, max_L, min(5, max_L), key=f"{attr}_L")
            with col2:
                U = st.slider("U", 0, max_U, min(5, max_U), key=f"{attr}_U")
            with col3:
                S = st.slider("S", 0, max_S, min(5, max_S), key=f"{attr}_S")
            with col4:
                R = st.slider("R", 0, max_R, min(5, max_R), key=f"{attr}_R")
            
            # Calculate ACS for this attribute using the formula
            acs_score = (0.35 * L + 0.30 * U + 0.20 * S + 0.15 * R)
            
            st.caption(f"**ACS = (0.35×{L} + 0.30×{U} + 0.20×{S} + 0.15×{R}) = {acs_score:.2f}**")
            
            acs_data.append({"Attribute": attr, "ACS": acs_score})
    
    # Calculate Total ACS and Weights
    st.markdown("---")
    st.header("Calculated Weights")
    
    total_acs = sum([item["ACS"] for item in acs_data])
    
    weight_results = []
    for item in acs_data:
        weight = (item["ACS"] / total_acs * 100) if total_acs > 0 else 0
        weight_results.append({
            "Attribute": item["Attribute"],
            "ACS Score": round(item["ACS"], 2),
            "Weight (%)": round(weight, 2)
        })
    
    df_weights = pd.DataFrame(weight_results)
    st.dataframe(df_weights, width='stretch', hide_index=True)
    
    # Display summary with explanation
    # Calculate the actual max possible ACS based on Table 6 max values
    max_possible_acs = 0
    for attr_data in lusr_table_6_data:
        max_L = attr_data.get("Legal Strength (L)", 10)
        max_U = attr_data.get("Uniqueness (U)", 10)
        max_S = attr_data.get("Stability (S)", 10)
        max_R = attr_data.get("Service Risk Impact (R)", 10)
        max_acs_attr = (0.35 * max_L + 0.30 * max_U + 0.20 * max_S + 0.15 * max_R)
        max_possible_acs += max_acs_attr
    
    col_summary1, col_summary2, col_summary3 = st.columns(3)
    with col_summary1:
        st.metric("Total ACS Score", f"{total_acs:.2f}/{max_possible_acs:.2f}", help=f"Actual range based on Table 6 max values: 0 to {max_possible_acs:.2f}")
    with col_summary2:
        st.metric("Total Weight %", "100.00", help="Always normalized to 100% regardless of ACS")
    with col_summary3:
        percentage = (total_acs / max_possible_acs * 100) if max_possible_acs > 0 else 0
        st.metric("% of Max ACS", f"{percentage:.1f}%", help=f"Your score as percentage of maximum possible")
    
    # Documentation below calculator
    st.markdown("---")
    st.header("Weight Calculation Documentation")
    
    with st.expander("Complete Weight Calculation Documentation", expanded=False):
        # Read the DOCX file
        doc_file = "data/Weight Calculation.docx"
        doc = Document(doc_file)
        
        # Display all paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Style headings
                if text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                    st.markdown(f"#### {text}")
                else:
                    st.markdown(text)
        
        # Display tables from document
        if doc.tables:
            st.markdown("### Reference Tables")
            for i, table in enumerate(doc.tables):
                st.markdown(f"#### Table {i+1}")
                
                # Convert table to dataframe
                data = []
                for row in table.rows:
                    data.append([cell.text for cell in row.cells])
                
                if data:
                    headers = data[0] if len(data) > 0 else []
                    rows = data[1:] if len(data) > 1 else []
                    df_table = pd.DataFrame(rows, columns=headers)
                    st.dataframe(df_table, width='stretch', hide_index=True)
    
    # Download documentation
    st.markdown("### Download Documentation")
    with open(doc_file, 'rb') as f:
        doc_bytes = f.read()
    st.download_button(
        label="Download DOCX",
        data=doc_bytes,
        file_name="Weight_Calculation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

elif page == "LUSR":
    st.markdown('<div class="info-card"><h1 style="color: white; margin: 0;">LUSR Calculation</h1><p style="color: rgba(255,255,255,0.9); margin-top: 0.5rem;">LUSR methodology and calculation framework</p></div>', unsafe_allow_html=True)
    
    # Read the DOCX file
    doc_file = "LUSR Calculation.docx"
    doc = Document(doc_file)
    
    # Extract and display document content
    st.markdown("### Complete LUSR Calculation Documentation")
    
    # Display all paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Style headings
            if text.startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
                st.markdown(f"#### {text}")
            else:
                st.markdown(text)
    
    # Display tables from document
    if doc.tables:
        st.markdown("### Reference Tables")
        for i, table in enumerate(doc.tables):
            st.markdown(f"#### Table {i+1}")
            
            # Convert table to dataframe
            data = []
            for row in table.rows:
                data.append([cell.text for cell in row.cells])
            
            if data:
                headers = data[0] if len(data) > 0 else []
                rows = data[1:] if len(data) > 1 else []
                df_table = pd.DataFrame(rows, columns=headers)
                st.dataframe(df_table, width='stretch', hide_index=True)
    
    # Add Table 6: LUSR Attribute Scoring Matrix
    st.markdown("#### Table 6: LUSR Attribute Scoring Matrix")
    
    lusr_table_6_data = [
        {"S.No": 1, "Attribute": "Aadhaar", "Legal Strength (L)": 10, "Uniqueness (U)": 10, "Stability (S)": 9, "Service Risk Impact (R)": 9, "ACL": 9.65},
        {"S.No": 2, "Attribute": "Name", "Legal Strength (L)": 7, "Uniqueness (U)": 5, "Stability (S)": 7, "Service Risk Impact (R)": 6, "ACL": 6.25},
        {"S.No": 3, "Attribute": "Date of Birth", "Legal Strength (L)": 8, "Uniqueness (U)": 6, "Stability (S)": 9, "Service Risk Impact (R)": 8, "ACL": 7.6},
        {"S.No": 4, "Attribute": "Mobile Number", "Legal Strength (L)": 6, "Uniqueness (U)": 6, "Stability (S)": 6, "Service Risk Impact (R)": 7, "ACL": 6.15},
        {"S.No": 5, "Attribute": "Gender", "Legal Strength (L)": 7, "Uniqueness (U)": 4, "Stability (S)": 9, "Service Risk Impact (R)": 6, "ACL": 6.35},
        {"S.No": 6, "Attribute": "Father's Name", "Legal Strength (L)": 8, "Uniqueness (U)": 5, "Stability (S)": 8, "Service Risk Impact (R)": 6, "ACL": 6.8},
        {"S.No": 7, "Attribute": "Mother's Name", "Legal Strength (L)": 8, "Uniqueness (U)": 5, "Stability (S)": 8, "Service Risk Impact (R)": 6, "ACL": 6.8},
        {"S.No": 8, "Attribute": "Permanent Address", "Legal Strength (L)": 7, "Uniqueness (U)": 5, "Stability (S)": 6, "Service Risk Impact (R)": 7, "ACL": 6.2},
        {"S.No": 9, "Attribute": "Correspondence Address", "Legal Strength (L)": 5, "Uniqueness (U)": 4, "Stability (S)": 5, "Service Risk Impact (R)": 5, "ACL": 4.7},
        {"S.No": 10, "Attribute": "Caste", "Legal Strength (L)": 8, "Uniqueness (U)": 5, "Stability (S)": 8, "Service Risk Impact (R)": 9, "ACL": 7.25},
        {"S.No": 11, "Attribute": "Marital Status", "Legal Strength (L)": 7, "Uniqueness (U)": 4, "Stability (S)": 6, "Service Risk Impact (R)": 6, "ACL": 5.75},
        {"S.No": 12, "Attribute": "Education Status", "Legal Strength (L)": 6, "Uniqueness (U)": 4, "Stability (S)": 7, "Service Risk Impact (R)": 5, "ACL": 5.45},
        {"S.No": 13, "Attribute": "Employment Status", "Legal Strength (L)": 6, "Uniqueness (U)": 4, "Stability (S)": 5, "Service Risk Impact (R)": 6, "ACL": 5.2},
        {"S.No": 14, "Attribute": "Ration Card Number", "Legal Strength (L)": 7, "Uniqueness (U)": 8, "Stability (S)": 7, "Service Risk Impact (R)": 8, "ACL": 7.45},
        {"S.No": 15, "Attribute": "Ration Card Type", "Legal Strength (L)": 6, "Uniqueness (U)": 4, "Stability (S)": 6, "Service Risk Impact (R)": 7, "ACL": 5.55},
        {"S.No": 16, "Attribute": "PAN ID", "Legal Strength (L)": 9, "Uniqueness (U)": 9, "Stability (S)": 9, "Service Risk Impact (R)": 9, "ACL": 9.0},
        {"S.No": 17, "Attribute": "Bank Account", "Legal Strength (L)": 8, "Uniqueness (U)": 8, "Stability (S)": 6, "Service Risk Impact (R)": 9, "ACL": 7.75},
        {"S.No": 18, "Attribute": "Land Ownership", "Legal Strength (L)": 9, "Uniqueness (U)": 7, "Stability (S)": 8, "Service Risk Impact (R)": 9, "ACL": 8.2},
        {"S.No": 19, "Attribute": "Motor Ownership", "Legal Strength (L)": 8, "Uniqueness (U)": 7, "Stability (S)": 7, "Service Risk Impact (R)": 7, "ACL": 7.35},
        {"S.No": 20, "Attribute": "Nationality", "Legal Strength (L)": 8, "Uniqueness (U)": 3, "Stability (S)": 9, "Service Risk Impact (R)": 5, "ACL": 6.25}
    ]
    
    df_table_6 = pd.DataFrame(lusr_table_6_data)
    st.dataframe(df_table_6, width='stretch', hide_index=True)
    
    # Add Table 7: LUSR Scoring Reference
    st.markdown("#### Table 7: LUSR Scoring Reference (Dimension-wise Scale)")
    
    lusr_table_7_data = [
        {"Dimension": "Legal Strength (L)", "Score Range": "10", "Condition Description": "Statutory Act backed + National Level Authority"},
        {"Dimension": "Legal Strength (L)", "Score Range": "8-9", "Condition Description": "State statutory registry or legal mandate"},
        {"Dimension": "Legal Strength (L)", "Score Range": "5-7", "Condition Description": "Official department database"},
        {"Dimension": "Legal Strength (L)", "Score Range": "3-4", "Condition Description": "Administrative / survey database"},
        {"Dimension": "Legal Strength (L)", "Score Range": "0-2", "Condition Description": "Self-declared / unverified"},
        {"Dimension": "Uniqueness (U)", "Score Range": "10", "Condition Description": "Biometric or globally unique identifier"},
        {"Dimension": "Uniqueness (U)", "Score Range": "8-9", "Condition Description": "System-generated unique ID"},
        {"Dimension": "Uniqueness (U)", "Score Range": "5-7", "Condition Description": "Combination-based uniqueness"},
        {"Dimension": "Uniqueness (U)", "Score Range": "<5", "Condition Description": "Common attribute with duplicates possible"},
        {"Dimension": "Stability (S)", "Score Range": "10", "Condition Description": "Never changes"},
        {"Dimension": "Stability (S)", "Score Range": "8-9", "Condition Description": "Rarely changes"},
        {"Dimension": "Stability (S)", "Score Range": "5-7", "Condition Description": "Occasionally changes"},
        {"Dimension": "Stability (S)", "Score Range": "<5", "Condition Description": "Frequently changes"},
        {"Dimension": "Service Risk Impact (R)", "Score Range": "10", "Condition Description": "Wrong value causes severe financial/legal impact"},
        {"Dimension": "Service Risk Impact (R)", "Score Range": "8-9", "Condition Description": "High scheme eligibility or DBT impact"},
        {"Dimension": "Service Risk Impact (R)", "Score Range": "5-7", "Condition Description": "Moderate service impact"},
        {"Dimension": "Service Risk Impact (R)", "Score Range": "<5", "Condition Description": "Low operational impact"}
    ]
    
    df_table_7 = pd.DataFrame(lusr_table_7_data)
    st.dataframe(df_table_7, width='stretch', hide_index=True)
    
    # Download documentation
    st.markdown("### Download Documentation")
    with open(doc_file, 'rb') as f:
        doc_bytes = f.read()
    st.download_button(
        label="Download DOCX",
        data=doc_bytes,
        file_name="LUSR_Calculation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )